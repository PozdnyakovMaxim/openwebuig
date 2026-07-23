from __future__ import annotations

from collections import Counter, defaultdict
from dataclasses import asdict, dataclass, field, replace
from io import BytesIO
from pathlib import Path
import hashlib
import json
import re
from typing import Any, Iterable
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from .ooxml_source import source_ooxml_inventory


DATE_RE = re.compile(r"\b\d{2}\.\d{2}\.\d{4}\b")
ORG_RE = re.compile(r"(?:ООО|АО|ПАО)\s+«[^»]+»", re.IGNORECASE)
INDEX_RE = re.compile(
    r"(?:ИНДЕКС|Индекс)\s+(?:НД|ЛНА)\s*[:：]?\s*(?P<value>[A-ZА-ЯЁ0-9./-]+)",
    re.IGNORECASE,
)
VERSION_RE = re.compile(r"ВЕРСИЯ\s*(?P<value>[\d.]+)", re.IGNORECASE)
EFFECTIVE_DATE_RE = re.compile(
    r"Введено\s+в\s+действие\s+с\s*[:：]?\s*(?P<date>\d{2}\.\d{2}\.\d{4})",
    re.IGNORECASE,
)
DECLARED_PAGES_RE = re.compile(r"Листов\s*[:：]?\s*(?P<count>\d+)", re.IGNORECASE)
APPROVAL_ORDER_RE = re.compile(
    r"от\s*(?P<date>\d{2}\.\d{2}\.\d{4})\s*№\s*(?P<number>.+)$",
    re.IGNORECASE,
)

TOC_TITLE_RE = re.compile(r"^СОДЕРЖАНИЕ$", re.IGNORECASE)
TOC_ENTRY_RE = re.compile(r"(?:\t|\s{2,}|\.+)\d+$|[А-Яа-яA-Za-zЁё)]\d+$")
APPENDIX_RE = re.compile(
    r"^Приложение\s*№?\s*(?P<number>\d+)(?:\s+(?P<title>.+))?$",
    re.IGNORECASE,
)

NUMERIC_PREFIX_RE = re.compile(r"^(?P<number>\d+(?:\.\d+)*)(?:[.)])?\s+(?P<text>.+)$")
LETTER_BULLET_RE = re.compile(r"^(?P<marker>[А-Яа-яЁёA-Za-z])\)\s+(?P<text>.+)$")
HEADING_STYLE_RE = re.compile(r"^(?:Heading|Заголовок)\s*(?P<level>\d+)?$", re.IGNORECASE)
TOC_STYLE_RE = re.compile(r"^(?:TOC|Оглавление|Содержание)", re.IGNORECASE)
DOCUMENT_KIND_RE = re.compile(
    r"^(?:ПОЛИТИКА|СТАНДАРТ|РЕГЛАМЕНТ|ИНСТРУКЦИЯ|ПОЛОЖЕНИЕ|ПОРЯДОК|ПРОЦЕДУРА)$",
    re.IGNORECASE,
)
INVENTORY_WORD_RE = re.compile(r"[0-9A-Za-zА-Яа-яЁё]+", re.UNICODE)
WORDPROCESSINGML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
MARKUP_COMPATIBILITY_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
RELATIONSHIPS_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
W_XML = f"{{{WORDPROCESSINGML_NS}}}"
MC_XML = f"{{{MARKUP_COMPATIBILITY_NS}}}"
R_XML = f"{{{RELATIONSHIPS_NS}}}"


@dataclass
class RawBlock:
    source_kind: str
    text: str
    style_name: str | None = None
    is_bold: bool = False
    numbering_label: str | None = None
    numbering_prefix: str | None = None
    numbering_level: int | None = None
    numbering_format: str | None = None


@dataclass(frozen=True)
class SourceStorySegment:
    part: str
    story: str
    text: str
    style: str = ""
    location: str = "paragraph"
    has_dynamic_page_field: bool = False


@dataclass(frozen=True)
class NumberingLevel:
    start: int = 1
    number_format: str = "decimal"
    level_text: str = "%1."
    suffix: str = "tab"
    restart_after_level: int | None = None
    is_legal: bool = False


class DocumentNumbering:
    def __init__(self, doc: DocumentObject) -> None:
        self._levels: dict[int, dict[int, NumberingLevel]] = {}
        self._num_to_abstract: dict[int, int] = {}
        self._overrides: dict[tuple[int, int], NumberingLevel] = {}
        self._counters: dict[int, dict[int, int]] = {}
        self._load(doc)

    def label_for(self, paragraph: Paragraph) -> tuple[str, str, int, str] | None:
        properties = self._paragraph_numbering(paragraph)
        if properties is None:
            return None
        num_id, level_index = properties
        abstract_id = self._num_to_abstract.get(num_id)
        if abstract_id is None:
            return None
        level = self._overrides.get((num_id, level_index)) or self._levels.get(abstract_id, {}).get(level_index)
        if level is None:
            return None

        counters = self._counters.setdefault(num_id, {})
        for deeper_level in [value for value in counters if value > level_index]:
            deeper = self._level_for(num_id, abstract_id, deeper_level)
            if deeper is None:
                del counters[deeper_level]
                continue
            restart_level = deeper.restart_after_level
            if restart_level == 0:
                continue
            trigger_level = deeper_level - 1 if restart_level is None else restart_level - 1
            if level_index <= trigger_level:
                del counters[deeper_level]
        for parent_level in range(level_index):
            if parent_level not in counters:
                parent = self._level_for(num_id, abstract_id, parent_level)
                counters[parent_level] = parent.start if parent else 1
        counters[level_index] = counters.get(level_index, level.start - 1) + 1

        label = re.sub(
            r"%(\d+)",
            lambda match: self._format_counter(
                counters.get(int(match.group(1)) - 1, 1),
                (
                    "decimal"
                    if level.is_legal
                    else self._format_for_level(num_id, abstract_id, int(match.group(1)) - 1)
                ),
            ),
            level.level_text,
        )
        label = _normalize_bullet_label(label).strip()
        if not label or level.number_format == "none":
            return None
        separator = "" if level.suffix == "nothing" else " "
        return label, f"{label}{separator}", level_index, level.number_format

    def _load(self, doc: DocumentObject) -> None:
        try:
            root = doc.part.numbering_part.element
        except (AttributeError, KeyError):
            return

        for abstract in root.findall(qn("w:abstractNum")):
            raw_id = abstract.get(qn("w:abstractNumId"))
            if raw_id is None:
                continue
            abstract_id = int(raw_id)
            levels: dict[int, NumberingLevel] = {}
            for level in abstract.findall(qn("w:lvl")):
                level_index = int(level.get(qn("w:ilvl")) or 0)
                levels[level_index] = self._read_level(level)
            self._levels[abstract_id] = levels

        for num in root.findall(qn("w:num")):
            raw_num_id = num.get(qn("w:numId"))
            abstract_ref = num.find(qn("w:abstractNumId"))
            if raw_num_id is None or abstract_ref is None:
                continue
            num_id = int(raw_num_id)
            abstract_id = int(abstract_ref.get(qn("w:val")) or 0)
            self._num_to_abstract[num_id] = abstract_id
            for override in num.findall(qn("w:lvlOverride")):
                level_index = int(override.get(qn("w:ilvl")) or 0)
                base = self._levels.get(abstract_id, {}).get(level_index, NumberingLevel())
                override_level = override.find(qn("w:lvl"))
                if override_level is not None:
                    base = self._read_level(override_level, fallback=base)
                start_override = override.find(qn("w:startOverride"))
                if start_override is not None:
                    base = replace(base, start=int(start_override.get(qn("w:val")) or base.start))
                self._overrides[(num_id, level_index)] = base

    @staticmethod
    def _read_level(level: Any, *, fallback: NumberingLevel | None = None) -> NumberingLevel:
        base = fallback or NumberingLevel()
        start = level.find(qn("w:start"))
        number_format = level.find(qn("w:numFmt"))
        level_text = level.find(qn("w:lvlText"))
        suffix = level.find(qn("w:suff"))
        restart = level.find(qn("w:lvlRestart"))
        is_legal = level.find(qn("w:isLgl"))
        return NumberingLevel(
            start=int(start.get(qn("w:val")) or base.start) if start is not None else base.start,
            number_format=(number_format.get(qn("w:val")) or base.number_format) if number_format is not None else base.number_format,
            level_text=(level_text.get(qn("w:val")) or base.level_text) if level_text is not None else base.level_text,
            suffix=(suffix.get(qn("w:val")) or base.suffix) if suffix is not None else base.suffix,
            restart_after_level=(
                int(restart.get(qn("w:val")) or 0)
                if restart is not None
                else base.restart_after_level
            ),
            is_legal=(
                _on_off_enabled(is_legal)
                if is_legal is not None
                else base.is_legal
            ),
        )

    @staticmethod
    def _paragraph_numbering(paragraph: Paragraph) -> tuple[int, int] | None:
        num_id: int | None = None
        level_index: int | None = None
        num_prs: list[Any] = []
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
            num_prs.append(paragraph._p.pPr.numPr)
        style = paragraph.style
        while style is not None:
            p_pr = style.element.pPr
            if p_pr is not None and p_pr.numPr is not None:
                num_prs.append(p_pr.numPr)
            style = style.base_style
        for num_pr in num_prs:
            if num_id is None and num_pr.numId is not None:
                num_id = int(num_pr.numId.val)
            if level_index is None and num_pr.ilvl is not None:
                level_index = int(num_pr.ilvl.val)
            if num_id is not None and level_index is not None:
                break
        if num_id is None:
            return None
        if num_id == 0:
            return None
        return num_id, level_index or 0

    def _level_for(self, num_id: int, abstract_id: int, level_index: int) -> NumberingLevel | None:
        return self._overrides.get((num_id, level_index)) or self._levels.get(abstract_id, {}).get(level_index)

    def _format_for_level(self, num_id: int, abstract_id: int, level_index: int) -> str:
        level = self._level_for(num_id, abstract_id, level_index)
        return level.number_format if level else "decimal"

    @staticmethod
    def _format_counter(value: int, number_format: str) -> str:
        if number_format in {"lowerLetter", "lowerLatin"}:
            return _alphabetic_number(value, "abcdefghijklmnopqrstuvwxyz")
        if number_format in {"upperLetter", "upperLatin"}:
            return _alphabetic_number(value, "ABCDEFGHIJKLMNOPQRSTUVWXYZ")
        if number_format == "russianLower":
            return _alphabetic_number(value, "абвгдежзиклмнопрстуфхцчшщэюя")
        if number_format == "russianUpper":
            return _alphabetic_number(value, "АБВГДЕЖЗИКЛМНОПРСТУФХЦЧШЩЭЮЯ")
        if number_format == "lowerRoman":
            return _roman_number(value).lower()
        if number_format == "upperRoman":
            return _roman_number(value)
        if number_format == "decimalZero":
            return f"{value:02d}"
        return str(value)


@dataclass
class DocumentMetadata:
    source_path: str
    source_name: str
    source_sha256: str
    doc_id: str
    index_code: str | None = None
    display_title: str | None = None
    document_kind: str | None = None
    organization: str | None = None
    version: str | None = None
    approval_date: str | None = None
    approval_order_number: str | None = None
    effective_date: str | None = None
    declared_pages: int | None = None
    title_lines: list[str] = field(default_factory=list)


@dataclass
class ContentBlock:
    block_id: str
    kind: str
    text: str
    source_kind: str
    section_number: str | None = None
    section_title: str | None = None
    subsection_number: str | None = None
    subsection_title: str | None = None
    item_number: str | None = None
    item_marker: str | None = None
    display_prefix: str | None = None
    heading_number: str | None = None
    appendix_number: str | None = None
    appendix_title: str | None = None
    heading_level: int | None = None
    section_path: list[str] = field(default_factory=list)
    section_labels: list[str] = field(default_factory=list)
    source_story: str | None = None
    source_parts: list[str] = field(default_factory=list)
    source_locations: list[str] = field(default_factory=list)
    source_occurrences: int | None = None


@dataclass
class StructuredDocument:
    metadata: DocumentMetadata
    blocks: list[ContentBlock]

    def to_dict(self) -> dict:
        return {
            "metadata": asdict(self.metadata),
            "block_count": len(self.blocks),
            "blocks": [asdict(block) for block in self.blocks],
        }


def _alphabetic_number(value: int, alphabet: str) -> str:
    if value <= 0 or not alphabet:
        return str(value)
    result: list[str] = []
    base = len(alphabet)
    current = value
    while current:
        current, remainder = divmod(current - 1, base)
        result.append(alphabet[remainder])
    return "".join(reversed(result))


def _normalize_bullet_label(value: str) -> str:
    translations = {
        "\uf0b7": "•",
        "\uf0a7": "▪",
        "\uf0d8": "➢",
    }
    return "".join(translations.get(character, character) for character in value)


def _on_off_enabled(element: Any) -> bool:
    value = element.get(qn("w:val"))
    return value is None or str(value).strip().lower() not in {"0", "false", "off", "no"}


def _roman_number(value: int) -> str:
    if value <= 0 or value >= 4000:
        return str(value)
    numerals = (
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    )
    result: list[str] = []
    remainder = value
    for number, numeral in numerals:
        while remainder >= number:
            result.append(numeral)
            remainder -= number
    return "".join(result)


def _ensure_unique_block_ids(blocks: list[ContentBlock]) -> None:
    seen: set[str] = set()
    for index, block in enumerate(blocks, start=1):
        base = block.block_id or f"block-{index}"
        candidate = base
        suffix = 2
        while candidate in seen:
            candidate = f"{base}--{suffix}"
            suffix += 1
        block.block_id = candidate
        seen.add(candidate)


def _normalize_text(text: str) -> str:
    cleaned = text.replace("\xa0", " ").replace("\r", "\n")
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{2,}", "\n", cleaned)
    return cleaned.strip()


def _classification_text(text: str) -> str:
    """Collapse Word line breaks before structural classification."""

    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()


def _paragraph_has_bold_text(paragraph: Paragraph) -> bool:
    bold_chars = 0
    total_chars = 0
    for run in paragraph.runs:
        text = run.text.strip()
        if not text:
            continue
        total_chars += len(text)
        if run.bold:
            bold_chars += len(text)
    return total_chars > 0 and bold_chars / total_chars >= 0.55


def _numbered_paragraph_text(
    paragraph: Paragraph,
    numbering: DocumentNumbering,
) -> tuple[str, str | None, str | None, int | None, str | None]:
    text = _normalize_text(paragraph.text)
    resolved = numbering.label_for(paragraph)
    if resolved is None:
        return text, None, None, None, None
    label, prefix, level_index, number_format = resolved
    if not text:
        return text, label, prefix, level_index, number_format

    compact_label = label.rstrip(".) ")
    already_numbered = (
        text.startswith(label)
        if prefix == label
        else text.startswith(prefix)
        or bool(
            compact_label
            and re.match(
                rf"^{re.escape(compact_label)}(?:[.)])?\s+",
                text,
                re.IGNORECASE,
            )
        )
    )
    if not already_numbered:
        text = f"{prefix}{text}"
    return text, label, prefix, level_index, number_format


def _strip_numbering_label(text: str, label: str | None) -> str:
    if not label:
        return text
    if text.startswith(label):
        return text[len(label) :].strip()
    return text


def _numbering_value(label: str | None) -> str | None:
    value = (label or "").strip()
    if not value:
        return None
    compact = re.sub(r"^[\s(\[]+", "", value)
    compact = re.sub(r"[.)\]\s]+$", "", compact).strip()
    return compact or value


def _iter_block_items(doc: DocumentObject) -> Iterable[RawBlock]:
    numbering = DocumentNumbering(doc)
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, doc)
            text, numbering_label, numbering_prefix, numbering_level, numbering_format = _numbered_paragraph_text(
                paragraph,
                numbering,
            )
            if text:
                yield RawBlock(
                    source_kind="paragraph",
                    text=text,
                    style_name=paragraph.style.name if paragraph.style else None,
                    is_bold=_paragraph_has_bold_text(paragraph),
                    numbering_label=numbering_label,
                    numbering_prefix=numbering_prefix,
                    numbering_level=numbering_level,
                    numbering_format=numbering_format,
                )
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for row in table.rows:
                parts: list[str] = []
                row_paragraphs: list[
                    tuple[str, str | None, str | None, int | None, str | None, Paragraph]
                ] = []
                seen_cells: set[int] = set()
                for cell in row.cells:
                    cell_key = id(cell._tc)
                    if cell_key in seen_cells:
                        continue
                    seen_cells.add(cell_key)
                    paragraphs = []
                    for paragraph in cell.paragraphs:
                        resolved = _numbered_paragraph_text(paragraph, numbering)
                        if resolved[0]:
                            row_paragraphs.append((*resolved, paragraph))
                            paragraphs.append(resolved[0])
                    cell_text = " ".join(paragraphs)
                    if cell_text:
                        parts.append(cell_text)
                text = " | ".join(parts)
                text = _normalize_text(text)
                if text:
                    single = row_paragraphs[0] if len(row_paragraphs) == 1 else None
                    yield RawBlock(
                        source_kind="table_row",
                        text=text,
                        style_name=(single[5].style.name if single and single[5].style else None),
                        is_bold=_paragraph_has_bold_text(single[5]) if single else False,
                        numbering_label=single[1] if single else None,
                        numbering_prefix=single[2] if single else None,
                        numbering_level=single[3] if single else None,
                        numbering_format=single[4] if single else None,
                    )


def _iter_header_footer_blocks(doc: DocumentObject) -> Iterable[RawBlock]:
    for section in doc.sections:
        for story in (section.header, section.footer):
            for paragraph in story.paragraphs:
                text = _normalize_text(paragraph.text)
                if text:
                    yield RawBlock(
                        source_kind="story_paragraph",
                        text=text,
                        style_name=paragraph.style.name if paragraph.style else None,
                        is_bold=_paragraph_has_bold_text(paragraph),
                    )
            for table in story.tables:
                for row in table.rows:
                    parts: list[str] = []
                    for cell in row.cells:
                        cell_text = " ".join(
                            _normalize_text(p.text)
                            for p in cell.paragraphs
                            if _normalize_text(p.text)
                        )
                        if cell_text:
                            parts.append(cell_text)
                    text = _normalize_text(" | ".join(parts))
                    if text:
                        yield RawBlock(source_kind="story_table_row", text=text)


def _source_story_inventory(doc: DocumentObject) -> list[SourceStorySegment]:
    segments = _story_paragraph_segments(
        doc.element,
        part_name="word/document.xml",
        story="body",
    )
    story_reference_ids = _visible_relationship_ids(
        doc.element,
        {f"{W_XML}headerReference", f"{W_XML}footerReference"},
    )
    note_ids = {
        "footnotes": _visible_note_ids(doc.element, f"{W_XML}footnoteReference"),
        "endnotes": _visible_note_ids(doc.element, f"{W_XML}endnoteReference"),
    }
    seen_parts: set[str] = set()
    relationships = sorted(doc.part.rels.values(), key=lambda relationship: relationship.rId)
    for relationship in relationships:
        relationship_type = str(relationship.reltype or "").rsplit("/", 1)[-1]
        if (
            relationship.is_external
            or relationship_type not in {"header", "footer", "footnotes", "endnotes"}
        ):
            continue
        if relationship_type in {"header", "footer"}:
            if relationship.rId not in story_reference_ids:
                continue
            note_filter = None
            story = relationship_type
        else:
            note_filter = note_ids[relationship_type]
            if not note_filter:
                continue
            story = relationship_type.removesuffix("s")

        target_part = relationship.target_part
        part_name = str(target_part.partname).lstrip("/")
        if part_name in seen_parts:
            continue
        seen_parts.add(part_name)
        root = getattr(target_part, "element", None)
        if root is None:
            root = parse_xml(target_part.blob)
        segments.extend(
            _story_paragraph_segments(
                root,
                part_name=part_name,
                story=story,
                note_ids=note_filter,
            )
        )
    return _filter_source_story_segments(segments)


def _story_paragraph_segments(
    root: Any,
    *,
    part_name: str,
    story: str,
    note_ids: set[str] | None = None,
) -> list[SourceStorySegment]:
    segments: list[SourceStorySegment] = []

    def visit(
        node: Any,
        *,
        in_textbox: bool = False,
        in_table: bool = False,
    ) -> None:
        if node.tag in {f"{W_XML}footnote", f"{W_XML}endnote"}:
            if node.get(f"{W_XML}type"):
                return
            if note_ids is not None and node.get(f"{W_XML}id") not in note_ids:
                return
        if node.tag == f"{MC_XML}AlternateContent":
            branch = _story_alternate_content_branch(node)
            if branch is not None:
                visit(branch, in_textbox=in_textbox, in_table=in_table)
            return

        nested_textbox = in_textbox or node.tag == f"{W_XML}txbxContent"
        nested_table = in_table or node.tag == f"{W_XML}tbl"
        if node.tag == f"{W_XML}p":
            text = _visible_story_paragraph_text(node)
            if text:
                segments.append(
                    SourceStorySegment(
                        part=part_name,
                        story=story,
                        text=text,
                        style=_story_paragraph_style(node),
                        location=(
                            "textbox"
                            if nested_textbox
                            else "table"
                            if nested_table
                            else "paragraph"
                        ),
                        has_dynamic_page_field=_story_has_dynamic_page_field(node),
                    )
                )
        for child in _story_visible_children(node):
            visit(child, in_textbox=nested_textbox, in_table=nested_table)

    visit(root)
    return segments


def _visible_story_paragraph_text(paragraph: Any) -> str:
    parts: list[str] = []

    def collect(node: Any, *, root: bool = False) -> None:
        if not root and node.tag == f"{W_XML}p":
            return
        if node.tag in {f"{W_XML}del", f"{W_XML}moveFrom"}:
            return
        if node.tag == f"{W_XML}r" and _story_run_is_hidden(node):
            return
        if node.tag == f"{W_XML}t":
            if node.text:
                parts.append(str(node.text))
            return
        if node.tag == f"{W_XML}tab":
            parts.append(" ")
            return
        if node.tag in {f"{W_XML}br", f"{W_XML}cr"}:
            parts.append("\n")
            return
        if node.tag == f"{MC_XML}AlternateContent":
            branch = _story_alternate_content_branch(node)
            if branch is not None:
                collect(branch)
            return
        for child in _story_visible_children(node):
            collect(child)

    collect(paragraph, root=True)
    return _normalize_text("".join(parts))


def _story_visible_children(node: Any) -> list[Any]:
    if node.tag != f"{MC_XML}AlternateContent":
        return list(node)
    branch = _story_alternate_content_branch(node)
    return [branch] if branch is not None else []


def _story_alternate_content_branch(node: Any) -> Any | None:
    choice = node.find(f"{MC_XML}Choice")
    return choice if choice is not None else node.find(f"{MC_XML}Fallback")


def _story_run_is_hidden(run: Any) -> bool:
    properties = run.find(f"{W_XML}rPr")
    if properties is None:
        return False
    return any(
        _story_boolean_is_on(properties.find(f"{W_XML}{name}"))
        for name in ("vanish", "webHidden")
    )


def _story_boolean_is_on(element: Any | None) -> bool:
    if element is None:
        return False
    value = str(element.get(f"{W_XML}val") or "true").strip().casefold()
    return value not in {"0", "false", "off", "no"}


def _story_paragraph_style(paragraph: Any) -> str:
    properties = paragraph.find(f"{W_XML}pPr")
    style = properties.find(f"{W_XML}pStyle") if properties is not None else None
    return str(style.get(f"{W_XML}val") or "") if style is not None else ""


def _story_has_dynamic_page_field(paragraph: Any) -> bool:
    instructions = " ".join(
        str(node.text or "") for node in paragraph.iter(f"{W_XML}instrText")
    ).upper()
    return bool(re.search(r"\b(?:PAGE|NUMPAGES|SECTIONPAGES)\b", instructions))


def _visible_relationship_ids(root: Any, tags: set[str]) -> set[str]:
    values: set[str] = set()

    def visit(node: Any) -> None:
        if node.tag in {f"{W_XML}del", f"{W_XML}moveFrom"}:
            return
        if node.tag == f"{W_XML}r" and _story_run_is_hidden(node):
            return
        if node.tag in tags:
            value = node.get(f"{R_XML}id")
            if value:
                values.add(str(value))
        for child in _story_visible_children(node):
            visit(child)

    visit(root)
    return values


def _visible_note_ids(root: Any, tag: str) -> set[str]:
    values: set[str] = set()

    def visit(node: Any) -> None:
        if node.tag in {f"{W_XML}del", f"{W_XML}moveFrom"}:
            return
        if node.tag == f"{W_XML}r" and _story_run_is_hidden(node):
            return
        if node.tag == tag:
            value = node.get(f"{W_XML}id")
            if value:
                values.add(str(value))
        for child in _story_visible_children(node):
            visit(child)

    visit(root)
    return values


def _filter_source_story_segments(
    segments: list[SourceStorySegment],
) -> list[SourceStorySegment]:
    required: list[SourceStorySegment] = []
    in_toc: dict[str, bool] = defaultdict(bool)
    for segment in segments:
        text = _normalize_text(segment.text)
        style = segment.style.casefold().replace(" ", "")
        toc_style = style.startswith(("toc", "contents", "оглавлен", "содержан"))
        if segment.story == "body":
            if text.casefold() in {"содержание", "оглавление"} or toc_style:
                in_toc[segment.part] = True
                continue
            if in_toc[segment.part] and _looks_like_toc_entry(text):
                continue
            in_toc[segment.part] = False
        if segment.story in {"header", "footer"} and _story_is_page_counter(
            segment,
            text,
        ):
            continue
        required.append(segment)
    return required


def _story_is_page_counter(segment: SourceStorySegment, text: str) -> bool:
    tokens = _inventory_tokens(text)
    if tokens and all(token.isdigit() for token in tokens):
        return True
    if re.fullmatch(
        r"(?:(?:стр(?:аница)?\.?|page)\s*(?:№\s*)?\d+"
        r"(?:\s*(?:из|of|/)\s*\d+)?|\d+\s*/\s*\d+)",
        text,
        re.IGNORECASE,
    ):
        return True
    if not segment.has_dynamic_page_field:
        return False
    page_words = {"стр", "страница", "page", "из", "of"}
    return bool(tokens) and all(
        token.isdigit()
        or token in page_words
        or bool(re.fullmatch(r"[ivxlcdm]+", token, re.IGNORECASE))
        for token in tokens
    )


def _inventory_tokens(text: str) -> list[str]:
    return [match.group(0).casefold() for match in INVENTORY_WORD_RE.finditer(text)]


def _inventory_subsequence_count(haystack: list[str], needle: list[str]) -> int:
    if not needle or len(needle) > len(haystack):
        return 0
    width = len(needle)
    return sum(
        haystack[index : index + width] == needle
        for index in range(len(haystack) - width + 1)
    )


def _primary_body_inventory_segments(
    blocks: list[ContentBlock],
) -> list[tuple[str, str]]:
    block_segments: list[tuple[str, str]] = []
    for block in blocks:
        if block.kind == "supplemental":
            continue
        text = _normalize_text(block.text)
        if not text:
            continue
        marker = ""
        # Automatic Word numbering is synthetic and absent from w:t; manual
        # prefixes remain part of the OOXML text and must be included.
        if not block.display_prefix:
            if block.kind == "heading":
                marker = block.heading_number or (
                    block.section_path[-1] if block.section_path else ""
                )
            elif block.kind in {"numbered_paragraph", "appendix_numbered_item"}:
                marker = block.item_number or ""
            elif block.kind in {"letter_bullet", "appendix_bullet", "list_item"}:
                marker = block.item_marker or ""
        rendered = _normalize_text(f"{marker} {text}")
        occurrences = max(1, int(block.source_occurrences or 1))
        block_segments.extend([(block.source_kind, rendered)] * occurrences)
    return block_segments


def _supplemental_content_blocks(
    _doc: DocumentObject,
    _metadata: DocumentMetadata,
    blocks: list[ContentBlock],
    *,
    source_bytes: bytes | None = None,
) -> list[ContentBlock]:
    if source_bytes is None:
        raise ValueError("source_bytes are required for canonical OOXML inventory")
    segments = source_ooxml_inventory(source_bytes)["segments"]
    grouped: dict[tuple[str, str], list[Any]] = defaultdict(list)
    for segment in segments:
        grouped[(segment.story, _normalize_text(segment.text))].append(segment)

    supplemental: list[ContentBlock] = []
    selected: set[tuple[str, str]] = set()

    def add_group(
        key: tuple[str, str],
        examples: list[Any],
        occurrences: int,
    ) -> None:
        story, text = key
        digest = hashlib.sha1(f"{story}\0{text}".encode("utf-8")).hexdigest()[:16]
        supplemental.append(
            ContentBlock(
                block_id=f"supplemental-{story}-{digest}",
                kind="supplemental",
                text=text,
                source_kind=f"supplemental_{story}",
                source_story=story,
                source_parts=sorted({item.part for item in examples}),
                source_locations=sorted({item.location for item in examples}),
                source_occurrences=occurrences,
            )
        )
        selected.add(key)

    for key in sorted(grouped):
        examples = grouped[key]
        if key[0] != "body":
            add_group(key, examples, len(examples))
            continue
        textbox_examples = [
            item for item in examples if item.location == "textbox"
        ]
        if textbox_examples:
            add_group(key, textbox_examples, len(textbox_examples))

    primary_body_segments = _primary_body_inventory_segments(blocks)
    primary_exact = Counter(
        tuple(_inventory_tokens(segment))
        for source_kind, segment in primary_body_segments
        if source_kind != "table_row"
    )
    primary_table_tokens = [
        _inventory_tokens(segment)
        for source_kind, segment in primary_body_segments
        if source_kind == "table_row"
    ]

    for key in sorted(grouped):
        if key in selected or key[0] != "body":
            continue
        examples = grouped[key]
        needle = _inventory_tokens(key[1])
        if not needle:
            continue
        actual_occurrences = primary_exact[tuple(needle)] + sum(
            _inventory_subsequence_count(tokens, needle)
            for tokens in primary_table_tokens
        )
        expected_occurrences = len(examples)
        missing_occurrences = max(0, expected_occurrences - actual_occurrences)
        if not missing_occurrences:
            continue
        add_group(key, examples, missing_occurrences)

    return supplemental


def _safe_slug(text: str) -> str:
    lowered = text.lower()
    slug = re.sub(r"[^a-zа-яё0-9]+", "-", lowered, flags=re.IGNORECASE)
    slug = slug.strip("-")
    return slug or "document"


def _looks_like_toc_entry(text: str) -> bool:
    candidate = _classification_text(text)
    if TOC_ENTRY_RE.search(candidate):
        return True
    if re.match(
        r"^(?:\d+(?:\.\d+)*|Приложение\s*№?\s*\d+).+\s+\d{1,4}$",
        candidate,
        re.IGNORECASE,
    ):
        return True
    return False


def _style_heading_level(block: RawBlock) -> int | None:
    if not block.style_name:
        return None
    match = HEADING_STYLE_RE.match(block.style_name.strip())
    if not match:
        return None
    raw_level = match.group("level")
    return int(raw_level) if raw_level else 1


def _style_is_toc(block: RawBlock) -> bool:
    return bool(block.style_name and TOC_STYLE_RE.match(block.style_name.strip()))


def _parse_numeric_prefix(text: str) -> tuple[str, str, int] | None:
    match = NUMERIC_PREFIX_RE.match(_classification_text(text))
    if not match:
        return None
    number = match.group("number")
    title = match.group("text").strip()
    if title.isdigit() or DATE_RE.fullmatch(number):
        return None
    return number, title, number.count(".") + 1


def _parse_level1(text: str) -> tuple[str, str] | None:
    parsed = _parse_numeric_prefix(text)
    if parsed and parsed[2] == 1:
        return parsed[0], parsed[1]
    return None


def _parse_level2(text: str) -> tuple[str, str] | None:
    parsed = _parse_numeric_prefix(text)
    if parsed and parsed[2] == 2:
        return parsed[0], parsed[1]
    return None


def _parse_numbered_paragraph(text: str) -> tuple[str, str] | None:
    parsed = _parse_numeric_prefix(text)
    if parsed and parsed[2] >= 3:
        return parsed[0], parsed[1]
    return None


def _parse_letter_bullet(text: str) -> tuple[str, str] | None:
    match = LETTER_BULLET_RE.match(_classification_text(text))
    if not match:
        return None
    return match.group("marker"), match.group("text").strip()


def _looks_like_content_start(text: str) -> bool:
    candidate = _classification_text(text)
    return bool(
        _parse_numeric_prefix(candidate)
        or APPENDIX_RE.match(candidate)
        or TOC_TITLE_RE.match(candidate)
    )


def _looks_like_heading(block: RawBlock, title: str, level: int) -> bool:
    if _style_heading_level(block):
        return True
    if block.source_kind == "table_row":
        return False
    if block.numbering_label and not block.is_bold:
        return False
    if level <= 2:
        return True
    if re.search(r"[.;:]$", title) or len(title) > 90:
        return False
    letters = [character for character in title if character.isalpha()]
    uppercase = bool(letters) and sum(character.isupper() for character in letters) / len(letters) >= 0.8
    return block.is_bold or uppercase


def _is_metadata_noise(text: str) -> bool:
    return bool(
        text.upper() == "УТВЕРЖДЕНО"
        or "приказом" in text.lower()
        or INDEX_RE.search(text)
        or EFFECTIVE_DATE_RE.search(text)
        or DECLARED_PAGES_RE.search(text)
        or VERSION_RE.search(text)
        or APPROVAL_ORDER_RE.search(text)
    )


def _extract_metadata(
    source_path: Path,
    raw_blocks: list[RawBlock],
    *,
    source_sha256: str,
) -> DocumentMetadata:
    index_code: str | None = None
    approval_date: str | None = None
    approval_order_number: str | None = None
    effective_date: str | None = None
    declared_pages: int | None = None
    version: str | None = None
    organization: str | None = None

    title_lines: list[str] = []
    started_title_capture = False

    meta_blocks = raw_blocks
    for block in meta_blocks:
        text = block.text

        if not index_code:
            match = INDEX_RE.search(text)
            if match:
                index_code = match.group("value").strip()
                started_title_capture = True
                continue

        if started_title_capture:
            if EFFECTIVE_DATE_RE.search(text) or DECLARED_PAGES_RE.search(text):
                started_title_capture = False
            elif TOC_TITLE_RE.match(text):
                started_title_capture = False
            elif _parse_level1(text) or _parse_level2(text) or _parse_numbered_paragraph(text):
                started_title_capture = False
            elif ORG_RE.fullmatch(text):
                organization = organization or text
            elif VERSION_RE.search(text):
                version = version or VERSION_RE.search(text).group("value")
            elif text.upper() == "УТВЕРЖДЕНО":
                continue
            elif "приказом" in text.lower():
                continue
            elif INDEX_RE.search(text):
                continue
            else:
                title_lines.append(text)

        if not approval_date or not approval_order_number:
            match = APPROVAL_ORDER_RE.search(text)
            if match:
                approval_date = approval_date or match.group("date")
                approval_order_number = approval_order_number or match.group("number").strip()

        if not effective_date:
            match = EFFECTIVE_DATE_RE.search(text)
            if match:
                effective_date = match.group("date")

        if declared_pages is None:
            match = DECLARED_PAGES_RE.search(text)
            if match:
                declared_pages = int(match.group("count"))

        if not version:
            match = VERSION_RE.search(text)
            if match:
                version = match.group("value")

        if not organization:
            match = ORG_RE.search(text)
            if match:
                organization = match.group(0)

    cleaned_title_lines = [line for line in title_lines if line]
    display_title = " ".join(cleaned_title_lines).strip() or None
    document_kind = cleaned_title_lines[0] if cleaned_title_lines else None

    if not version and index_code:
        fallback = re.search(r"-(\d+(?:\.\d+)*)$", index_code)
        if fallback:
            version = fallback.group(1)

    doc_id_source = index_code or display_title or source_path.stem
    return DocumentMetadata(
        source_path=str(source_path),
        source_name=source_path.name,
        source_sha256=source_sha256,
        doc_id=_safe_slug(doc_id_source),
        index_code=index_code,
        display_title=display_title,
        document_kind=document_kind,
        organization=organization,
        version=version,
        approval_date=approval_date,
        approval_order_number=approval_order_number,
        effective_date=effective_date,
        declared_pages=declared_pages,
        title_lines=cleaned_title_lines,
    )


def extract_docx(
    source_path: str | Path,
    *,
    _source_bytes: bytes | None = None,
) -> StructuredDocument:
    path = Path(source_path)
    source_bytes = _source_bytes if _source_bytes is not None else path.read_bytes()
    source_sha256 = hashlib.sha256(source_bytes).hexdigest()
    doc = Document(BytesIO(source_bytes))
    header_footer_blocks = list(_iter_header_footer_blocks(doc))
    body_blocks = list(_iter_block_items(doc))
    metadata = _extract_metadata(
        path,
        header_footer_blocks + body_blocks,
        source_sha256=source_sha256,
    )

    blocks: list[ContentBlock] = []
    hierarchy: list[tuple[int, str | None, str]] = []
    current_appendix_number: str | None = None
    current_appendix_title: str | None = None
    content_started = False
    in_toc = False
    pending_appendix_title = False

    def context() -> dict:
        hierarchy_by_level = {level: (level, number, title) for level, number, title in hierarchy}
        section = hierarchy_by_level.get(1)
        subsection = hierarchy_by_level.get(2)
        return {
            "section_number": section[1] if section else None,
            "section_title": section[2] if section else None,
            "subsection_number": subsection[1] if subsection else None,
            "subsection_title": subsection[2] if subsection else None,
            "section_path": [number for _, number, _ in hierarchy if number],
            "section_labels": [
                " ".join(part for part in (number, title) if part).strip()
                for _, number, title in hierarchy
            ],
        }

    def append_heading(raw_block: RawBlock, number: str | None, title: str, level: int) -> None:
        nonlocal content_started, hierarchy
        hierarchy = [item for item in hierarchy if item[0] < level]
        hierarchy.append((level, number, title))
        content_started = True
        ctx = context()
        block_prefix = number or f"unnumbered-{len(blocks)+1}"
        blocks.append(
            ContentBlock(
                block_id=f"heading-{block_prefix}",
                kind="heading",
                text=title,
                source_kind=raw_block.source_kind,
                display_prefix=raw_block.numbering_prefix if raw_block.numbering_label else None,
                heading_number=number,
                heading_level=level,
                **ctx,
            )
        )

    for raw_block in body_blocks:
        text = raw_block.text

        if TOC_TITLE_RE.match(text) or _style_is_toc(raw_block):
            in_toc = True
            continue

        if in_toc:
            if _looks_like_toc_entry(text) or _style_is_toc(raw_block):
                continue
            in_toc = False

        appendix_match = APPENDIX_RE.match(_classification_text(text))
        if appendix_match:
            current_appendix_number = appendix_match.group("number")
            inline_appendix_title = appendix_match.group("title")
            current_appendix_title = inline_appendix_title.strip() if inline_appendix_title else None
            hierarchy = []
            content_started = True
            pending_appendix_title = current_appendix_title is None
            blocks.append(
                ContentBlock(
                    block_id=f"appendix-{current_appendix_number}",
                    kind="appendix_heading",
                    text=text,
                    source_kind=raw_block.source_kind,
                    appendix_number=current_appendix_number,
                    appendix_title=current_appendix_title,
                )
            )
            continue

        if pending_appendix_title:
            current_appendix_title = text
            pending_appendix_title = False
            blocks.append(
                ContentBlock(
                    block_id=f"appendix-{current_appendix_number}-title",
                    kind="appendix_title",
                    text=text,
                    source_kind=raw_block.source_kind,
                    appendix_number=current_appendix_number,
                    appendix_title=current_appendix_title,
                )
            )
            continue

        if current_appendix_number is not None:
            if raw_block.numbering_format == "bullet" and raw_block.numbering_label:
                blocks.append(
                    ContentBlock(
                        block_id=f"appendix-{current_appendix_number}-{len(blocks)+1}",
                        kind="appendix_bullet",
                        text=_strip_numbering_label(text, raw_block.numbering_label),
                        source_kind=raw_block.source_kind,
                        item_marker=raw_block.numbering_label,
                        display_prefix=raw_block.numbering_prefix,
                        appendix_number=current_appendix_number,
                        appendix_title=current_appendix_title,
                    )
                )
                continue

            bullet = _parse_letter_bullet(text)
            if bullet:
                marker, body = bullet
                blocks.append(
                    ContentBlock(
                        block_id=f"appendix-{current_appendix_number}-{len(blocks)+1}",
                        kind="appendix_bullet",
                        text=body,
                        source_kind=raw_block.source_kind,
                        item_marker=marker,
                        appendix_number=current_appendix_number,
                        appendix_title=current_appendix_title,
                    )
                )
                continue

            automatic_number = _numbering_value(raw_block.numbering_label)
            if automatic_number:
                blocks.append(
                    ContentBlock(
                        block_id=f"appendix-{current_appendix_number}-item-{automatic_number}",
                        kind="appendix_numbered_item",
                        text=_strip_numbering_label(text, raw_block.numbering_label),
                        source_kind=raw_block.source_kind,
                        item_number=automatic_number,
                        display_prefix=raw_block.numbering_prefix,
                        appendix_number=current_appendix_number,
                        appendix_title=current_appendix_title,
                    )
                )
            elif numbered := _parse_numeric_prefix(text):
                number, body, _level = numbered
                blocks.append(
                    ContentBlock(
                        block_id=f"appendix-{current_appendix_number}-item-{number}",
                        kind="appendix_numbered_item",
                        text=body,
                        source_kind=raw_block.source_kind,
                        item_number=number,
                        appendix_number=current_appendix_number,
                        appendix_title=current_appendix_title,
                    )
                )
            else:
                blocks.append(
                    ContentBlock(
                        block_id=f"appendix-{current_appendix_number}-{len(blocks)+1}",
                        kind="appendix_paragraph",
                        text=text,
                        source_kind=raw_block.source_kind,
                        appendix_number=current_appendix_number,
                        appendix_title=current_appendix_title,
                    )
                )
            continue

        style_heading_level = _style_heading_level(raw_block)
        numbered_prefix = _parse_numeric_prefix(text)

        if style_heading_level:
            automatic_number = _numbering_value(raw_block.numbering_label)
            if automatic_number and raw_block.numbering_format != "bullet":
                append_heading(
                    raw_block,
                    automatic_number,
                    _strip_numbering_label(text, raw_block.numbering_label),
                    style_heading_level,
                )
            elif numbered_prefix:
                number, title, numeric_level = numbered_prefix
                append_heading(raw_block, number, title, numeric_level)
            else:
                append_heading(raw_block, None, text, style_heading_level)
            continue

        if raw_block.numbering_format == "bullet" and raw_block.numbering_label:
            content_started = True
            blocks.append(
                ContentBlock(
                    block_id=f"bullet-{len(blocks)+1}",
                    kind="list_item",
                    text=_strip_numbering_label(text, raw_block.numbering_label),
                    source_kind=raw_block.source_kind,
                    item_marker=raw_block.numbering_label,
                    display_prefix=raw_block.numbering_prefix,
                    **context(),
                )
            )
            continue

        automatic_number = _numbering_value(raw_block.numbering_label)
        if automatic_number:
            title = _strip_numbering_label(text, raw_block.numbering_label)
            numeric_level = (raw_block.numbering_level or 0) + 1
            if _looks_like_heading(raw_block, title, numeric_level):
                append_heading(raw_block, automatic_number, title, numeric_level)
            else:
                content_started = True
                blocks.append(
                    ContentBlock(
                        block_id=f"item-{automatic_number}",
                        kind="numbered_paragraph",
                        text=title,
                        source_kind=raw_block.source_kind,
                        item_number=automatic_number,
                        display_prefix=raw_block.numbering_prefix,
                        **context(),
                    )
                )
            continue

        if numbered_prefix:
            number, body, numeric_level = numbered_prefix
            if _looks_like_heading(raw_block, body, numeric_level):
                append_heading(raw_block, number, body, numeric_level)
                continue
            content_started = True
            blocks.append(
                ContentBlock(
                    block_id=f"item-{number}",
                    kind="numbered_paragraph",
                    text=body,
                    source_kind=raw_block.source_kind,
                    item_number=number,
                    **context(),
                )
            )
            continue

        bullet = _parse_letter_bullet(text)
        if bullet and content_started:
            marker, body = bullet
            blocks.append(
                ContentBlock(
                    block_id=f"bullet-{len(blocks)+1}",
                    kind="letter_bullet",
                    text=body,
                    source_kind=raw_block.source_kind,
                    item_marker=marker,
                    **context(),
                )
            )
            continue

        if content_started:
            blocks.append(
                ContentBlock(
                    block_id=f"paragraph-{len(blocks)+1}",
                    kind="paragraph",
                    text=text,
                    source_kind=raw_block.source_kind,
                    **context(),
                )
            )
        elif not _is_metadata_noise(text):
            blocks.append(
                ContentBlock(
                    block_id=f"front-matter-{len(blocks)+1}",
                    kind="front_matter",
                    text=text,
                    source_kind=raw_block.source_kind,
                )
            )

    blocks.extend(
        _supplemental_content_blocks(
            doc,
            metadata,
            blocks,
            source_bytes=source_bytes,
        )
    )
    _ensure_unique_block_ids(blocks)
    return StructuredDocument(metadata=metadata, blocks=blocks)


def extract_docx_text(
    source_path: str | Path,
    *,
    expected_sha256: str | None = None,
) -> str:
    path = Path(source_path)
    source_bytes = path.read_bytes()
    actual_sha256 = hashlib.sha256(source_bytes).hexdigest()
    if expected_sha256 and actual_sha256 != expected_sha256.strip().casefold():
        raise ValueError("Source DOCX changed since it was indexed")
    doc = Document(BytesIO(source_bytes))
    body_texts = [block.text for block in _iter_block_items(doc) if block.text]
    seen = {_normalize_text(text) for text in body_texts}
    supplemental_texts: list[str] = []
    for block in extract_docx(path, _source_bytes=source_bytes).blocks:
        if block.kind != "supplemental":
            continue
        text = _normalize_text(block.text)
        if not text or text in seen:
            continue
        seen.add(text)
        supplemental_texts.append(text)
    return "\n\n".join(body_texts + supplemental_texts).strip()


def extract_many(paths: Iterable[str | Path]) -> list[StructuredDocument]:
    return [extract_docx(path) for path in paths]


def write_extraction(doc: StructuredDocument, destination: str | Path) -> Path:
    output_path = Path(destination)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(
        json.dumps(doc.to_dict(), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return output_path
