from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory
import unittest
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document

from document_search.ooxml_source import source_ooxml_inventory


class OoxmlSourceInventoryTest(unittest.TestCase):
    def test_inventory_reads_visible_vml_textbox_and_body_numbers(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "vml-textbox.docx"
            document = Document()
            document.add_paragraph("1.4 Основные принципы")
            document.save(source_path)
            _inject_document_fragment(
                source_path,
                (
                    '<w:p><w:r><w:pict>'
                    '<v:shape xmlns:v="urn:schemas-microsoft-com:vml">'
                    "<v:textbox><w:txbxContent><w:p><w:r>"
                    "<w:t>КОММЕРЧЕСКАЯ ТАЙНА</w:t>"
                    "</w:r></w:p></w:txbxContent></v:textbox>"
                    "</v:shape></w:pict></w:r></w:p>"
                ),
            )

            inventory = source_ooxml_inventory(source_path)

        observed = {
            (segment.story, segment.location, segment.text)
            for segment in inventory["segments"]
        }
        self.assertIn(("body", "paragraph", "1.4 Основные принципы"), observed)
        self.assertIn(("body", "textbox", "КОММЕРЧЕСКАЯ ТАЙНА"), observed)

    def test_inventory_reads_drawingml_textbox_from_bytes(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "drawing-textbox.docx"
            Document().save(source_path)
            _inject_document_fragment(
                source_path,
                (
                    '<w:p><w:r><w:drawing>'
                    '<wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
                    "<wps:txbx><w:txbxContent><w:p><w:r>"
                    "<w:t>Текст DrawingML</w:t>"
                    "</w:r></w:p></w:txbxContent></wps:txbx>"
                    "</wps:wsp></w:drawing></w:r></w:p>"
                ),
            )

            inventory = source_ooxml_inventory(source_path.read_bytes())

        matching = [
            segment
            for segment in inventory["segments"]
            if segment.text == "Текст DrawingML"
        ]
        self.assertEqual(len(matching), 1)
        self.assertEqual(matching[0].location, "textbox")

    def test_alternate_content_uses_one_branch_and_excludes_hidden_text(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "alternate-content.docx"
            Document().save(source_path)
            _inject_document_fragment(
                source_path,
                (
                    '<mc:AlternateContent '
                    'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">'
                    '<mc:Choice Requires="wps"><w:p><w:r>'
                    "<w:t>Видимая ветка</w:t>"
                    "</w:r></w:p></mc:Choice>"
                    '<mc:Fallback><w:p><w:r><w:t>Резервная ветка</w:t>'
                    "</w:r></w:p></mc:Fallback></mc:AlternateContent>"
                    '<w:p><w:r><w:rPr><w:vanish/></w:rPr>'
                    "<w:t>Скрытый текст</w:t></w:r></w:p>"
                ),
            )

            inventory = source_ooxml_inventory(source_path)

        texts = [segment.text for segment in inventory["segments"]]
        self.assertEqual(texts.count("Видимая ветка"), 1)
        self.assertNotIn("Резервная ветка", texts)
        self.assertNotIn("Скрытый текст", texts)


def _inject_document_fragment(path: Path, fragment: str) -> None:
    with ZipFile(path) as source:
        files = {item.filename: source.read(item.filename) for item in source.infolist()}
    document_xml = files["word/document.xml"].decode("utf-8")
    document_xml = document_xml.replace("<w:sectPr", f"{fragment}<w:sectPr", 1)
    files["word/document.xml"] = document_xml.encode("utf-8")

    temporary_path = path.with_suffix(".rewrite.docx")
    with ZipFile(temporary_path, "w", ZIP_DEFLATED) as destination:
        for name, payload in files.items():
            destination.writestr(name, payload)
    temporary_path.replace(path)


if __name__ == "__main__":
    unittest.main()
