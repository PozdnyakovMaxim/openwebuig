from __future__ import annotations

from io import BytesIO
import unittest
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document

from document_search.external_document_loader import (
    DocumentTooLargeError,
    InvalidDocumentError,
    UnsupportedDocumentError,
    _validate_docx,
    process_external_document,
)


class ExternalDocumentLoaderTest(unittest.TestCase):
    def test_docx_automatic_numbering_is_present_in_openwebui_response(self) -> None:
        document = Document()
        document.add_paragraph("Первое автоматическое требование", style="List Number")
        document.add_paragraph("Второе автоматическое требование", style="List Number")
        buffer = BytesIO()
        document.save(buffer)

        result = process_external_document(
            buffer.getvalue(),
            filename="%D0%9F%D0%BE%D0%BB%D0%B8%D1%82%D0%B8%D0%BA%D0%B0.docx",
            content_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )

        text = str(result["page_content"])
        self.assertIn("1. Первое автоматическое требование", text)
        self.assertIn("2. Второе автоматическое требование", text)
        metadata = result["metadata"]
        assert isinstance(metadata, dict)
        self.assertEqual(metadata["file_name"], "Политика.docx")
        self.assertTrue(metadata["word_numbering_preserved"])

    def test_rejects_invalid_docx_package(self) -> None:
        with self.assertRaises(InvalidDocumentError):
            process_external_document(
                b"not-a-zip",
                filename="broken.docx",
                content_type="application/octet-stream",
            )

    def test_rejects_malformed_ooxml_as_invalid_document(self) -> None:
        buffer = BytesIO()
        with ZipFile(buffer, "w") as archive:
            archive.writestr("[Content_Types].xml", "<broken")
            archive.writestr("word/document.xml", "<broken")

        with self.assertRaisesRegex(InvalidDocumentError, "could not be parsed"):
            process_external_document(
                buffer.getvalue(),
                filename="broken.docx",
            )

    def test_rejects_suspicious_docx_compression_ratio(self) -> None:
        buffer = BytesIO()
        with ZipFile(buffer, "w", compression=ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", "<Types/>")
            archive.writestr("word/document.xml", b"0" * (2 * 1024 * 1024))

        with self.assertRaisesRegex(InvalidDocumentError, "compression ratio"):
            process_external_document(
                buffer.getvalue(),
                filename="bomb.docx",
            )

    def test_rejects_docx_member_count_and_uncompressed_size_limits(self) -> None:
        buffer = BytesIO()
        with ZipFile(buffer, "w") as archive:
            archive.writestr("[Content_Types].xml", "<Types/>")
            archive.writestr("word/document.xml", "<document/>")
            archive.writestr("word/extra.xml", "extra")
        payload = buffer.getvalue()

        with self.assertRaisesRegex(InvalidDocumentError, "too many ZIP members"):
            _validate_docx(payload, max_member_count=2)
        with self.assertRaisesRegex(InvalidDocumentError, "uncompressed size"):
            _validate_docx(payload, max_uncompressed_bytes=10)

    def test_rejects_unsupported_binary_format(self) -> None:
        with self.assertRaises(UnsupportedDocumentError):
            process_external_document(
                b"binary",
                filename="sheet.xlsx",
                content_type="application/octet-stream",
            )

    def test_enforces_size_limit_before_parsing(self) -> None:
        with self.assertRaises(DocumentTooLargeError):
            process_external_document(
                b"12345",
                filename="notes.txt",
                content_type="text/plain",
                max_bytes=4,
            )


if __name__ == "__main__":
    unittest.main()
