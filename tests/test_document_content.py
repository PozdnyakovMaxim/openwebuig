from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory
import hashlib
import unittest

from docx import Document

from document_search.document_content import find_source_document, load_source_document_text


class DocumentContentTest(unittest.TestCase):
    def test_finds_nested_source_file_and_extracts_full_body(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            source_path = root / "AIDIT" / "Политика.docx"
            source_path.parent.mkdir()
            document = Document()
            document.add_paragraph("Первый абзац")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Ячейка 1"
            table.cell(0, 1).text = "Ячейка 2"
            document.add_paragraph("Последний абзац")
            document.save(source_path)

            record = {"source_name": "Политика.docx", "metadata": {}}
            found = find_source_document(record, roots=[root])

            self.assertEqual(found, source_path)
            text = load_source_document_text(
                {
                    "source_name": "Политика.docx",
                    "metadata": {"source_path": str(source_path)},
                }
            )
            self.assertEqual(
                text,
                "Первый абзац\n\nЯчейка 1 | Ячейка 2\n\nПоследний абзац",
            )

    def test_changed_source_file_is_not_returned_as_indexed_full_text(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "Политика.docx"
            document = Document()
            document.add_paragraph("Исходная редакция")
            document.save(source_path)
            original_hash = hashlib.sha256(source_path.read_bytes()).hexdigest()

            changed = Document()
            changed.add_paragraph("Новая неиндексированная редакция")
            changed.save(source_path)

            text = load_source_document_text(
                {
                    "source_name": source_path.name,
                    "metadata": {
                        "source_path": str(source_path),
                        "source_sha256": original_hash,
                    },
                }
            )

        self.assertIsNone(text)


if __name__ == "__main__":
    unittest.main()
