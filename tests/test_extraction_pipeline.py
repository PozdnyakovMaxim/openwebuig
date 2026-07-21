from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any
import unittest
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from document_search.chunker import chunk_document
from document_search.extractor import extract_docx, extract_docx_text


def _value_element(tag: str, value: object) -> Any:
    element = OxmlElement(tag)
    element.set(qn("w:val"), str(value))
    return element


def _add_numbering_definition(
    document: Any,
    levels: list[dict[str, object]],
    *,
    start_overrides: dict[int, int] | None = None,
) -> int:
    root = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")) or 0)
        for element in root.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(element.get(qn("w:numId")) or 0)
        for element in root.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    abstract.append(_value_element("w:multiLevelType", "multilevel"))
    for level_index, values in enumerate(levels):
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), str(level_index))
        level.append(_value_element("w:start", values.get("start", 1)))
        level.append(_value_element("w:numFmt", values.get("format", "decimal")))
        level.append(_value_element("w:lvlText", values.get("text", f"%{level_index + 1}.")))
        level.append(_value_element("w:suff", values.get("suffix", "space")))
        if "restart" in values:
            level.append(_value_element("w:lvlRestart", values["restart"]))
        if values.get("legal"):
            level.append(OxmlElement("w:isLgl"))
        abstract.append(level)

    first_num_index = next(
        (index for index, child in enumerate(root) if child.tag == qn("w:num")),
        len(root),
    )
    root.insert(first_num_index, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    num.append(_value_element("w:abstractNumId", abstract_id))
    for level_index, start in sorted((start_overrides or {}).items()):
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), str(level_index))
        override.append(_value_element("w:startOverride", start))
        num.append(override)
    root.append(num)
    return num_id


def _append_num_pr(
    properties: Any,
    *,
    num_id: int | None,
    level: int | None,
) -> None:
    existing = properties.find(qn("w:numPr"))
    if existing is not None:
        properties.remove(existing)
    num_pr = OxmlElement("w:numPr")
    if level is not None:
        num_pr.append(_value_element("w:ilvl", level))
    if num_id is not None:
        num_pr.append(_value_element("w:numId", num_id))
    properties.append(num_pr)


def _number_paragraph(paragraph: Any, num_id: int | None, level: int | None = 0) -> None:
    _append_num_pr(paragraph._p.get_or_add_pPr(), num_id=num_id, level=level)


def _number_style(style: Any, num_id: int, level: int | None = 0) -> None:
    _append_num_pr(style.element.get_or_add_pPr(), num_id=num_id, level=level)


def _content_chunks(chunked: dict[str, Any]) -> list[dict[str, Any]]:
    return [
        chunk
        for chunk in chunked["chunks"]
        if chunk.get("chunk_type") != "metadata"
    ]


class ExtractionPipelineTest(unittest.TestCase):
    def test_first_paragraph_after_toc_is_retained(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "toc.docx"
            document = Document()
            document.add_paragraph("СОДЕРЖАНИЕ")
            document.add_paragraph("1 Общие положения ........ 3")
            document.add_paragraph("После содержания этот абзац должен сохраниться.")
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()

        texts = [block["text"] for block in extracted["blocks"]]
        self.assertIn("После содержания этот абзац должен сохраниться.", texts)
        self.assertRegex(extracted["metadata"]["source_sha256"], r"^[0-9a-f]{64}$")

    def test_multiline_toc_entries_do_not_poison_following_structure(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "multiline-toc.docx"
            document = Document()
            document.add_paragraph("СОДЕРЖАНИЕ")
            document.add_paragraph("2.2\nТребования к резервному копированию\n5")
            document.add_paragraph("Приложение\n№ 3 Формы контроля\n17")
            document.add_paragraph(
                "2.2\nТребования к резервному копированию",
                style="Heading 2",
            )
            document.add_paragraph("Основной текст после многострочного оглавления.")
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()

        headings = [
            block
            for block in extracted["blocks"]
            if block["kind"] == "heading" and block.get("heading_number") == "2.2"
        ]
        self.assertEqual(len(headings), 1)
        self.assertEqual(headings[0]["text"], "Требования к резервному копированию")
        self.assertFalse(
            any(block["kind"] == "appendix_heading" for block in extracted["blocks"])
        )
        body = next(
            block
            for block in extracted["blocks"]
            if block["text"] == "Основной текст после многострочного оглавления."
        )
        self.assertEqual(body["section_path"], ["2.2"])
        self.assertIsNone(body["appendix_number"])

    def test_repeated_numbered_items_receive_unique_block_ids(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "duplicates.docx"
            document = Document()
            document.add_paragraph("1.1.1 Первый пункт.")
            document.add_paragraph("1.1.1 Второй пункт.")
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()

        block_ids = [block["block_id"] for block in extracted["blocks"]]
        self.assertEqual(len(block_ids), len(set(block_ids)))
        self.assertTrue(any(block_id.endswith("--2") for block_id in block_ids))

    def test_every_substantive_block_is_present_in_chunks(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "coverage.docx"
            document = Document()
            document.add_paragraph("ИНДЕКС НД: TEST-1.0")
            document.add_paragraph("ПОЛИТИКА")
            document.add_paragraph("1 Общие положения")
            document.add_paragraph("1.1 Назначение")
            document.add_paragraph("1.1.1 Первый обязательный пункт.")
            document.add_paragraph("Обычный поясняющий абзац.")
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()
            chunked = chunk_document(extracted)

        expected = {
            block["block_id"]
            for block in extracted["blocks"]
            if str(block.get("text") or "").strip()
        }
        covered = {
            block_id
            for chunk in chunked["chunks"]
            for block_id in chunk.get("block_ids") or []
            if block_id != "metadata"
        }
        self.assertEqual(expected, covered)

    def test_direct_numpr_override_and_suffix_survive_every_representation(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "direct-numbering.docx"
            document = Document()
            num_id = _add_numbering_definition(
                document,
                [{"format": "decimal", "text": "%1)", "suffix": "space"}],
                start_overrides={0: 4},
            )
            paragraph = document.add_paragraph("Прямой автоматический пункт")
            _number_paragraph(paragraph, num_id)
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()
            chunked = chunk_document(extracted)
            full_text = extract_docx_text(source_path)

        item = next(block for block in extracted["blocks"] if block["kind"] == "numbered_paragraph")
        self.assertEqual(item["item_number"], "4")
        self.assertEqual(item["display_prefix"], "4) ")
        self.assertEqual(item["text"], "Прямой автоматический пункт")
        chunk = next(chunk for chunk in _content_chunks(chunked) if "Прямой" in chunk["raw_text"])
        self.assertEqual(chunk["raw_text"], "4) Прямой автоматический пункт")
        self.assertIn("4) Прямой автоматический пункт", chunk["searchable_text"])
        self.assertIn("пункт 4", chunk["citation_label"])
        self.assertEqual(full_text, "4) Прямой автоматический пункт")

    def test_visible_prefix_is_not_duplicated_or_confused_with_content_digits(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "visible-prefix.docx"
            document = Document()
            num_id = _add_numbering_definition(
                document,
                [{"format": "decimal", "text": "%1", "suffix": "space"}],
            )
            first = document.add_paragraph("100 сотрудников участвуют")
            second = document.add_paragraph("2 Уже содержит отображаемый номер")
            _number_paragraph(first, num_id)
            _number_paragraph(second, num_id)
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()
            chunked = chunk_document(extracted)
            full_text = extract_docx_text(source_path)

        items = [block for block in extracted["blocks"] if block["kind"] == "numbered_paragraph"]
        self.assertEqual([block["text"] for block in items], ["100 сотрудников участвуют", "Уже содержит отображаемый номер"])
        raw_text = "\n".join(chunk["raw_text"] for chunk in _content_chunks(chunked))
        self.assertIn("1 100 сотрудников участвуют", raw_text)
        self.assertIn("2 Уже содержит отображаемый номер", raw_text)
        self.assertNotIn("2 2 Уже", raw_text)
        self.assertIn("1 100 сотрудников участвуют", full_text)

    def test_style_inheritance_merges_direct_level_with_base_style_num_id(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "style-inheritance.docx"
            document = Document()
            num_id = _add_numbering_definition(
                document,
                [
                    {"format": "decimal", "text": "%1."},
                    {"format": "decimal", "text": "%1.%2."},
                ],
            )
            base_style = document.styles.add_style("Numbered Base", WD_STYLE_TYPE.PARAGRAPH)
            _number_style(base_style, num_id, level=0)
            derived_style = document.styles.add_style("Numbered Derived", WD_STYLE_TYPE.PARAGRAPH)
            derived_style.base_style = base_style
            first = document.add_paragraph("Первый унаследованный подпункт", style=derived_style)
            second = document.add_paragraph("Второй унаследованный подпункт", style=derived_style)
            _number_paragraph(first, None, level=1)
            _number_paragraph(second, None, level=1)
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()
            chunked = chunk_document(extracted)

        items = [block for block in extracted["blocks"] if block["kind"] == "numbered_paragraph"]
        self.assertEqual([block["item_number"] for block in items], ["1.1", "1.2"])
        raw_text = "\n".join(chunk["raw_text"] for chunk in _content_chunks(chunked))
        self.assertIn("1.1. Первый унаследованный подпункт", raw_text)
        self.assertIn("1.2. Второй унаследованный подпункт", raw_text)

    def test_multilevel_numbering_restarts_after_parent_by_default(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "multilevel-restart.docx"
            document = Document()
            num_id = _add_numbering_definition(
                document,
                [
                    {"format": "decimal", "text": "%1."},
                    {"format": "decimal", "text": "%1.%2."},
                ],
            )
            for text, level in (
                ("Первый раздел", 0),
                ("Первый подпункт", 1),
                ("Второй подпункт", 1),
                ("Второй раздел", 0),
                ("Подпункт после перезапуска", 1),
            ):
                paragraph = document.add_paragraph(text)
                _number_paragraph(paragraph, num_id, level)
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()
            full_text = extract_docx_text(source_path)

        items = [block for block in extracted["blocks"] if block["kind"] == "numbered_paragraph"]
        self.assertEqual(
            [block["item_number"] for block in items],
            ["1", "1.1", "1.2", "2", "2.1"],
        )
        self.assertIn("2.1. Подпункт после перезапуска", full_text)

    def test_lvl_restart_zero_and_is_lgl_are_honored(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "legal-numbering.docx"
            document = Document()
            num_id = _add_numbering_definition(
                document,
                [
                    {"format": "upperRoman", "text": "%1."},
                    {
                        "format": "lowerLetter",
                        "text": "%1.%2)",
                        "restart": 0,
                        "legal": True,
                    },
                ],
            )
            for text, level in (
                ("Римский первый", 0),
                ("Юридический подпункт", 1),
                ("Римский второй", 0),
                ("Продолженный юридический подпункт", 1),
            ):
                paragraph = document.add_paragraph(text)
                _number_paragraph(paragraph, num_id, level)
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()
            chunked = chunk_document(extracted)
            full_text = extract_docx_text(source_path)

        items = [block for block in extracted["blocks"] if block["kind"] == "numbered_paragraph"]
        self.assertEqual([block["item_number"] for block in items], ["I", "1.1", "II", "2.2"])
        self.assertIn("I. Римский первый", full_text)
        self.assertIn("2.2) Продолженный юридический подпункт", full_text)
        roman_chunk = next(chunk for chunk in _content_chunks(chunked) if "Римский второй" in chunk["raw_text"])
        self.assertIn("II. Римский второй", roman_chunk["raw_text"])
        self.assertIn("пункт II", roman_chunk["citation_label"])

    def test_lvl_restart_uses_one_based_trigger_level(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "explicit-level-restart.docx"
            document = Document()
            num_id = _add_numbering_definition(
                document,
                [
                    {"format": "decimal", "text": "%1."},
                    {"format": "decimal", "text": "%1.%2."},
                    {"format": "decimal", "text": "%1.%2.%3.", "restart": 1},
                ],
            )
            for text, level in (
                ("Первый верхний уровень", 0),
                ("Первый средний уровень", 1),
                ("Первый нижний уровень", 2),
                ("Второй средний уровень", 1),
                ("Нижний уровень без перезапуска", 2),
                ("Второй верхний уровень", 0),
                ("Средний после перезапуска", 1),
                ("Нижний после перезапуска", 2),
            ):
                paragraph = document.add_paragraph(text)
                _number_paragraph(paragraph, num_id, level)
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()

        items = [block for block in extracted["blocks"] if block["kind"] == "numbered_paragraph"]
        self.assertEqual(
            [block["item_number"] for block in items],
            ["1", "1.1", "1.1.1", "1.2", "1.2.2", "2", "2.1", "2.1.1"],
        )

    def test_letter_numbering_and_nothing_suffix_are_preserved(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "letter-numbering.docx"
            document = Document()
            num_id = _add_numbering_definition(
                document,
                [{"start": 27, "format": "lowerLetter", "text": "(%1)", "suffix": "nothing"}],
            )
            paragraph = document.add_paragraph("Буквенный пункт без пробела")
            _number_paragraph(paragraph, num_id)
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()
            chunked = chunk_document(extracted)
            full_text = extract_docx_text(source_path)

        item = next(block for block in extracted["blocks"] if block["kind"] == "numbered_paragraph")
        self.assertEqual(item["item_number"], "aa")
        self.assertEqual(item["display_prefix"], "(aa)")
        chunk = next(chunk for chunk in _content_chunks(chunked) if "Буквенный" in chunk["raw_text"])
        self.assertEqual(chunk["raw_text"], "(aa)Буквенный пункт без пробела")
        self.assertIn("пункт aa", chunk["citation_label"])
        self.assertEqual(full_text, "(aa)Буквенный пункт без пробела")

    def test_automatic_bullets_keep_their_glyph_without_added_parenthesis(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "automatic-bullets.docx"
            document = Document()
            bullet_id = _add_numbering_definition(
                document,
                [{"format": "bullet", "text": "\uf0b7", "suffix": "space"}],
            )
            letter_bullet_id = _add_numbering_definition(
                document,
                [{"format": "bullet", "text": "o", "suffix": "space"}],
            )
            first = document.add_paragraph("Круглый маркер")
            second = document.add_paragraph("Буквоподобный маркер")
            _number_paragraph(first, bullet_id)
            _number_paragraph(second, letter_bullet_id)
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()
            chunked = chunk_document(extracted)
            full_text = extract_docx_text(source_path)

        bullets = [block for block in extracted["blocks"] if block["kind"] == "list_item"]
        self.assertEqual([block["item_marker"] for block in bullets], ["•", "o"])
        raw_text = "\n".join(chunk["raw_text"] for chunk in _content_chunks(chunked))
        self.assertIn("• Круглый маркер", raw_text)
        self.assertIn("o Буквоподобный маркер", raw_text)
        self.assertNotIn("o) Буквоподобный маркер", raw_text)
        self.assertIn("• Круглый маркер", full_text)

    def test_auto_numbered_heading_uses_heading_style_for_hierarchy_level(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "automatic-heading.docx"
            document = Document()
            num_id = _add_numbering_definition(
                document,
                [{"format": "decimal", "text": "%1."}],
            )
            heading = document.add_paragraph("Автоматический заголовок", style="Heading 2")
            _number_paragraph(heading, num_id, level=0)
            document.add_paragraph("Текст под заголовком")
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()
            chunked = chunk_document(extracted)
            full_text = extract_docx_text(source_path)

        heading_block = next(block for block in extracted["blocks"] if block["kind"] == "heading")
        self.assertEqual(heading_block["heading_number"], "1")
        self.assertEqual(heading_block["heading_level"], 2)
        self.assertIsNone(heading_block["section_number"])
        self.assertEqual(heading_block["subsection_number"], "1")
        self.assertEqual(heading_block["section_path"], ["1"])
        heading_chunk = next(chunk for chunk in _content_chunks(chunked) if chunk["chunk_type"] == "heading")
        self.assertEqual(heading_chunk["raw_text"], "1. Автоматический заголовок")
        self.assertIn("Раздел: 1 Автоматический заголовок", heading_chunk["searchable_text"])
        self.assertIn("раздел 1 Автоматический заголовок", heading_chunk["citation_label"])
        self.assertIn("1. Автоматический заголовок", full_text)

    def test_bold_deep_numbered_clause_is_not_misclassified_as_heading(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "deep-clause.docx"
            document = Document()
            document.add_paragraph("3 Общие положения", style="Heading 1")
            document.add_paragraph("3.4 Ответственные лица", style="Heading 2")
            clause = document.add_paragraph()
            clause.add_run(
                "3.4.5 Куратором настоящего ЛНА, ответственным за его актуализацию, "
                "является начальник отдела поддержки ИТ инфраструктуры."
            ).bold = True
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()

        item = next(block for block in extracted["blocks"] if block.get("item_number") == "3.4.5")
        self.assertEqual(item["kind"], "numbered_paragraph")
        self.assertEqual(item["section_path"], ["3", "3.4"])
        self.assertFalse(
            any(block.get("heading_number") == "3.4.5" for block in extracted["blocks"])
        )

    def test_word_automatic_numbering_survives_extraction_and_chunking(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "automatic-numbering.docx"
            document = Document()
            document.add_paragraph("Первый автоматический пункт", style="List Number")
            document.add_paragraph("Второй автоматический пункт", style="List Number")
            document.add_paragraph("Автоматический маркер", style="List Bullet")
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()
            chunked = chunk_document(extracted)
            full_text = extract_docx_text(source_path)

        numbered = [
            block
            for block in extracted["blocks"]
            if block.get("kind") == "numbered_paragraph"
        ]
        self.assertEqual([block["item_number"] for block in numbered], ["1", "2"])
        combined_chunks = "\n".join(chunk["raw_text"] for chunk in chunked["chunks"])
        self.assertIn("1. Первый автоматический пункт", combined_chunks)
        self.assertIn("2. Второй автоматический пункт", combined_chunks)
        self.assertIn("• Автоматический маркер", combined_chunks)
        self.assertIn("1. Первый автоматический пункт", full_text)
        self.assertIn("2. Второй автоматический пункт", full_text)

    def test_word_numbering_inside_table_survives_extraction_chunks_and_full_text(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "numbered-table.docx"
            document = Document()
            num_id = _add_numbering_definition(
                document,
                [{"format": "decimal", "text": "%1.", "suffix": "space"}],
            )
            table = document.add_table(rows=1, cols=1)
            paragraph = table.cell(0, 0).paragraphs[0]
            paragraph.add_run("Пункт внутри таблицы")
            _number_paragraph(paragraph, num_id)
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()
            chunked = chunk_document(extracted)
            full_text = extract_docx_text(source_path)

        item = next(block for block in extracted["blocks"] if block["kind"] == "numbered_paragraph")
        self.assertEqual(item["source_kind"], "table_row")
        self.assertEqual(item["item_number"], "1")
        chunk = next(chunk for chunk in _content_chunks(chunked) if "внутри таблицы" in chunk["raw_text"])
        self.assertEqual(chunk["raw_text"], "1. Пункт внутри таблицы")
        self.assertIn("1. Пункт внутри таблицы", chunk["searchable_text"])
        self.assertIn("пункт 1", chunk["citation_label"])
        self.assertIn("1. Пункт внутри таблицы", full_text)

    def test_supplemental_header_is_deduplicated_and_page_counter_is_ignored(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "header-footer.docx"
            document = Document()
            document.add_paragraph("Основной текст документа.")
            header = document.sections[0].header
            header.paragraphs[0].text = "Конфиденциальный служебный заголовок"
            header.add_paragraph("Конфиденциальный служебный заголовок")
            header.add_paragraph("Конфиденциальный служебный заголовок")
            document.sections[0].footer.paragraphs[0].text = "Стр. 1 из 1"
            document.save(source_path)

            extracted = extract_docx(source_path).to_dict()
            chunked = chunk_document(extracted)
            full_text = extract_docx_text(source_path)

        supplemental = [
            block for block in extracted["blocks"] if block["kind"] == "supplemental"
        ]
        self.assertEqual(len(supplemental), 1)
        self.assertEqual(supplemental[0]["source_story"], "header")
        self.assertEqual(supplemental[0]["source_occurrences"], 3)
        self.assertNotIn("Стр. 1 из 1", [block["text"] for block in supplemental])
        supplemental_chunks = [
            chunk for chunk in chunked["chunks"] if chunk["chunk_type"] == "supplemental"
        ]
        self.assertEqual(len(supplemental_chunks), 1)
        self.assertEqual(
            supplemental_chunks[0]["raw_text"].count("Конфиденциальный служебный заголовок"),
            1,
        )
        self.assertIn("верхний колонтитул", supplemental_chunks[0]["citation_label"])
        self.assertEqual(full_text.count("Конфиденциальный служебный заголовок"), 1)
        self.assertNotIn("Стр. 1 из 1", full_text)

    def test_textbox_and_referenced_notes_become_supplemental_searchable_blocks(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "supplemental-stories.docx"
            document = Document()
            document.add_paragraph("Основной текст документа.")
            document.save(source_path)
            _inject_supplemental_stories(source_path)

            extracted = extract_docx(source_path).to_dict()
            chunked = chunk_document(extracted)
            full_text = extract_docx_text(source_path)

        supplemental = [
            block for block in extracted["blocks"] if block["kind"] == "supplemental"
        ]
        by_story = {block["source_story"]: block for block in supplemental}
        self.assertEqual(set(by_story), {"body", "footnote", "endnote"})
        self.assertEqual(by_story["body"]["source_locations"], ["textbox"])
        for text in (
            "Текст внутри текстового поля.",
            "Текст связанной сноски.",
            "Текст связанной концевой сноски.",
        ):
            self.assertEqual(full_text.count(text), 1)
            self.assertEqual(
                sum(chunk["raw_text"].count(text) for chunk in chunked["chunks"]),
                1,
            )

    def test_skipped_body_content_control_cannot_be_masked_by_metadata_tokens(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source_path = Path(temporary_directory) / "body-content-control.docx"
            document = Document()
            document.add_paragraph("ИНДЕКС НД: TEST-1.0")
            document.save(source_path)
            _inject_body_content_control(source_path, "ИНДЕКС")

            extracted = extract_docx(source_path).to_dict()
            chunked = chunk_document(extracted)
            full_text = extract_docx_text(source_path)

        recovered = [
            block
            for block in extracted["blocks"]
            if block["kind"] == "supplemental"
            and block["source_story"] == "body"
            and block["text"] == "ИНДЕКС"
        ]
        self.assertEqual(len(recovered), 1)
        self.assertEqual(recovered[0]["source_locations"], ["paragraph"])
        self.assertTrue(
            any(
                chunk["chunk_type"] == "supplemental"
                and "ИНДЕКС" in chunk["raw_text"].splitlines()
                for chunk in chunked["chunks"]
            )
        )
        self.assertIn("ИНДЕКС", full_text)


def _inject_supplemental_stories(path: Path) -> None:
    with ZipFile(path) as source:
        files = {item.filename: source.read(item.filename) for item in source.infolist()}

    content_types = files["[Content_Types].xml"].decode("utf-8")
    content_types = content_types.replace(
        "</Types>",
        (
            '<Override PartName="/word/footnotes.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
            '<Override PartName="/word/endnotes.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"/>'
            "</Types>"
        ),
    )
    files["[Content_Types].xml"] = content_types.encode("utf-8")

    relationships_path = "word/_rels/document.xml.rels"
    relationships = files[relationships_path].decode("utf-8")
    relationships = relationships.replace(
        "</Relationships>",
        (
            '<Relationship Id="rIdAuditFootnotes" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" '
            'Target="footnotes.xml"/>'
            '<Relationship Id="rIdAuditEndnotes" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes" '
            'Target="endnotes.xml"/>'
            "</Relationships>"
        ),
    )
    files[relationships_path] = relationships.encode("utf-8")

    document_xml = files["word/document.xml"].decode("utf-8")
    references = (
        '<w:r><w:footnoteReference w:id="1"/></w:r>'
        '<w:r><w:endnoteReference w:id="1"/></w:r>'
    )
    document_xml = document_xml.replace("</w:p>", f"{references}</w:p>", 1)
    textbox = (
        '<w:p><w:r><w:pict><v:shape xmlns:v="urn:schemas-microsoft-com:vml">'
        '<v:textbox><w:txbxContent><w:p><w:r><w:t>'
        "Текст внутри текстового поля."
        "</w:t></w:r></w:p></w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>"
    )
    document_xml = document_xml.replace("<w:sectPr", f"{textbox}<w:sectPr", 1)
    files["word/document.xml"] = document_xml.encode("utf-8")

    files["word/footnotes.xml"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:footnote w:id="1"><w:p><w:r><w:t>Текст связанной сноски.</w:t>'
        "</w:r></w:p></w:footnote></w:footnotes>"
    ).encode("utf-8")
    files["word/endnotes.xml"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:endnote w:id="1"><w:p><w:r><w:t>Текст связанной концевой сноски.</w:t>'
        "</w:r></w:p></w:endnote></w:endnotes>"
    ).encode("utf-8")

    temporary_path = path.with_suffix(".rewrite.docx")
    with ZipFile(temporary_path, "w", ZIP_DEFLATED) as destination:
        for name, payload in files.items():
            destination.writestr(name, payload)
    temporary_path.replace(path)


def _inject_body_content_control(path: Path, text: str) -> None:
    with ZipFile(path) as source:
        files = {item.filename: source.read(item.filename) for item in source.infolist()}

    document_xml = files["word/document.xml"].decode("utf-8")
    content_control = (
        "<w:sdt><w:sdtContent><w:p><w:r><w:t>"
        f"{text}"
        "</w:t></w:r></w:p></w:sdtContent></w:sdt>"
    )
    document_xml = document_xml.replace("<w:sectPr", f"{content_control}<w:sectPr", 1)
    files["word/document.xml"] = document_xml.encode("utf-8")

    temporary_path = path.with_suffix(".rewrite.docx")
    with ZipFile(temporary_path, "w", ZIP_DEFLATED) as destination:
        for name, payload in files.items():
            destination.writestr(name, payload)
    temporary_path.replace(path)


if __name__ == "__main__":
    unittest.main()
