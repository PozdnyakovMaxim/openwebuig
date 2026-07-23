from __future__ import annotations

import json
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest
from unittest.mock import patch
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document

from document_search.chunker import chunk_document
from document_search import corpus_audit as corpus_audit_module
from document_search.corpus_audit import audit_corpus
from document_search.extractor import extract_docx, write_extraction


class CorpusAuditTest(unittest.TestCase):
    def test_audit_delegates_to_shared_ooxml_inventory(self) -> None:
        expected = {
            "segments": [],
            "ignored_segments": [],
            "story_counts": {},
            "location_counts": {},
        }

        with patch(
            "document_search.corpus_audit.source_ooxml_inventory",
            return_value=expected,
        ) as inventory:
            actual = corpus_audit_module._source_ooxml_inventory(Path("sample.docx"))

        self.assertIs(actual, expected)
        inventory.assert_called_once_with(Path("sample.docx"))

    def test_plain_duplicate_does_not_create_a_false_missing_textbox_error(self) -> None:
        text = "КОММЕРЧЕСКАЯ ТАЙНА"
        inventory = {
            "segments": [
                corpus_audit_module.SourceTextSegment(
                    part="word/document.xml",
                    story="body",
                    text=text,
                    location="paragraph",
                ),
                corpus_audit_module.SourceTextSegment(
                    part="word/document.xml",
                    story="body",
                    text=text,
                    location="textbox",
                ),
            ],
            "ignored_segments": [],
            "story_counts": {"body": 2},
            "location_counts": {"paragraph": 1, "textbox": 1},
        }
        extracted = {
            "metadata": {},
            "blocks": [
                {
                    "kind": "supplemental",
                    "text": text,
                    "source_story": "body",
                    "source_locations": ["textbox"],
                    "source_occurrences": 1,
                }
            ],
        }

        metrics = corpus_audit_module._source_inventory_metrics(inventory, extracted)

        self.assertEqual(metrics["source_missing_tokens"], 2)
        self.assertEqual(metrics["source_actionable_missing_segments"], 0)
        self.assertEqual(metrics["source_critical_missing_segments"], 0)

    def test_clean_artifacts_pass(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            paths = self._build_corpus(Path(temporary_directory))
            report = audit_corpus(*paths, skip_database=True)

        self.assertEqual(report["status"], "ok", report["issues"])
        self.assertEqual(report["summary"]["source_documents"], 1)
        self.assertEqual(report["summary"]["errors"], 0)
        self.assertEqual(
            report["documents"][0]["covered_blocks"],
            report["documents"][0]["indexable_blocks"],
        )
        self.assertEqual(report["documents"][0]["chunk_text_coverage"], 1.0)
        self.assertEqual(report["documents"][0]["chunk_snapshot_changed"], 0)
        self.assertEqual(report["summary"]["unsearchable_non_indexed_blocks"], 0)
        self.assertEqual(report["summary"]["source_missing_tokens"], 0)
        self.assertEqual(report["summary"]["source_token_coverage"], 1.0)

    def test_independent_ooxml_inventory_detects_repeated_footnote_text(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            paths = self._build_corpus(Path(temporary_directory))
            source_path = paths[0] / "Политика.docx"
            _inject_footnote(
                source_path,
                "Первый обязательный пункт документа.",
            )

            report = audit_corpus(*paths, skip_database=True)

        matching = [
            issue
            for issue in report["issues"]
            if issue["code"] == "source_ooxml_text_missing"
        ]
        self.assertEqual(report["status"], "error")
        self.assertEqual(len(matching), 1)
        self.assertEqual(matching[0]["level"], "error")
        self.assertGreater(report["documents"][0]["source_missing_tokens"], 0)
        self.assertGreater(report["documents"][0]["source_critical_missing_segments"], 0)
        self.assertIn("footnote", report["documents"][0]["source_story_counts"])

    def test_unreferenced_ooxml_story_does_not_create_false_positive(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            paths = self._build_corpus(Path(temporary_directory))
            _inject_footnote(
                paths[0] / "Политика.docx",
                "Несвязанная старая сноска внутри архива.",
                add_reference=False,
            )

            report = audit_corpus(*paths, skip_database=True)

        self.assertNotIn(
            "source_ooxml_text_missing",
            {issue["code"] for issue in report["issues"]},
        )
        self.assertNotIn("footnote", report["documents"][0]["source_story_counts"])

    def test_literal_page_counter_is_ignored_in_headers_and_footers(self) -> None:
        segments = [
            corpus_audit_module.SourceTextSegment(
                part="word/footer1.xml",
                story="footer",
                text="Стр. 2 из 7",
            ),
            corpus_audit_module.SourceTextSegment(
                part="word/header1.xml",
                story="header",
                text="Политика резервного копирования",
            ),
            corpus_audit_module.SourceTextSegment(
                part="word/header1.xml",
                story="header",
                text="Политика — стр. 2 из 7",
                has_dynamic_page_field=True,
            ),
        ]

        required, ignored = corpus_audit_module._partition_source_segments(segments)

        self.assertEqual(
            [item.text for item in required],
            ["Политика резервного копирования", "Политика — стр. 2 из 7"],
        )
        self.assertEqual([item.text for item in ignored], ["Стр. 2 из 7"])

    def test_missing_block_is_reported(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            paths = self._build_corpus(Path(temporary_directory))
            chunks_path = next(paths[2].glob("*.chunks.json"))
            data = json.loads(chunks_path.read_text(encoding="utf-8"))
            data["chunks"] = [
                chunk
                for chunk in data["chunks"]
                if "item-1.1.2" not in (chunk.get("block_ids") or [])
            ]
            data["chunk_count"] = len(data["chunks"])
            chunks_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            manifest_path = paths[2] / "manifest.json"
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            manifest[0]["chunks"] = len(data["chunks"])
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            report = audit_corpus(*paths, skip_database=True)

        codes = {issue["code"] for issue in report["issues"]}
        self.assertEqual(report["status"], "error")
        self.assertIn("chunk_blocks_missing", codes)
        self.assertIn("chunk_text_coverage_low", codes)

    def test_stale_extraction_is_reported(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            paths = self._build_corpus(Path(temporary_directory))
            extraction_path = next(paths[1].glob("*.json"))
            if extraction_path.name == "manifest.json":
                extraction_path = next(
                    path for path in paths[1].glob("*.json") if path.name != "manifest.json"
                )
            data = json.loads(extraction_path.read_text(encoding="utf-8"))
            data["blocks"] = [
                block for block in data["blocks"] if block.get("block_id") != "item-1.1.2"
            ]
            data["block_count"] = len(data["blocks"])
            extraction_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

            report = audit_corpus(*paths, skip_database=True)

        codes = {issue["code"] for issue in report["issues"]}
        self.assertEqual(report["status"], "error")
        self.assertIn("extraction_blocks_missing", codes)

    def test_source_sha256_mismatch_is_reported(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            paths = self._build_corpus(Path(temporary_directory))
            extraction_path = next(
                path for path in paths[1].glob("*.json") if path.name != "manifest.json"
            )
            extraction = json.loads(extraction_path.read_text(encoding="utf-8"))
            extraction["metadata"]["source_sha256"] = "0" * 64
            extraction_path.write_text(
                json.dumps(extraction, ensure_ascii=False),
                encoding="utf-8",
            )

            report = audit_corpus(*paths, skip_database=True)

        issue = next(
            issue for issue in report["issues"] if issue["code"] == "extraction_metadata_stale"
        )
        self.assertEqual(report["status"], "error")
        self.assertIn("source_sha256", issue["details"]["fields"])

    def test_empty_source_directory_is_reported(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            docs_dir = root / "docs"
            extracted_dir = root / "extracted"
            chunks_dir = root / "chunks"
            for directory in (docs_dir, extracted_dir, chunks_dir):
                directory.mkdir()
            for directory in (extracted_dir, chunks_dir):
                (directory / "manifest.json").write_text("[]", encoding="utf-8")

            report = audit_corpus(
                docs_dir,
                extracted_dir,
                chunks_dir,
                skip_database=True,
            )

        codes = {issue["code"] for issue in report["issues"]}
        self.assertEqual(report["status"], "error")
        self.assertIn("source_documents_missing", codes)

    def test_missing_metadata_chunk_is_reported(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            paths = self._build_corpus(Path(temporary_directory))
            chunks_path = next(paths[2].glob("*.chunks.json"))
            data = json.loads(chunks_path.read_text(encoding="utf-8"))
            data["chunks"] = [
                chunk for chunk in data["chunks"] if chunk.get("chunk_type") != "metadata"
            ]
            data["chunk_count"] = len(data["chunks"])
            chunks_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            manifest_path = paths[2] / "manifest.json"
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            manifest[0]["chunks"] = len(data["chunks"])
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            report = audit_corpus(*paths, skip_database=True)

        codes = {issue["code"] for issue in report["issues"]}
        self.assertEqual(report["status"], "error")
        self.assertIn("metadata_chunk_count_invalid", codes)

    def test_changed_searchable_text_is_reported(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            paths = self._build_corpus(Path(temporary_directory))
            chunks_path = next(paths[2].glob("*.chunks.json"))
            data = json.loads(chunks_path.read_text(encoding="utf-8"))
            data["chunks"][-1]["searchable_text"] += " устаревшее значение"
            chunks_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

            report = audit_corpus(*paths, skip_database=True)

        codes = {issue["code"] for issue in report["issues"]}
        self.assertEqual(report["status"], "error")
        self.assertIn("chunk_snapshot_changed", codes)

    def test_heading_missing_from_chunks_is_reported(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            paths = self._build_corpus(Path(temporary_directory))
            extraction_path = next(
                path for path in paths[1].glob("*.json") if path.name != "manifest.json"
            )
            extraction = json.loads(extraction_path.read_text(encoding="utf-8"))
            heading_id = next(
                block["block_id"]
                for block in extraction["blocks"]
                if block.get("kind") == "heading" and block.get("text") == "Назначение"
            )
            chunks_path = next(paths[2].glob("*.chunks.json"))
            data = json.loads(chunks_path.read_text(encoding="utf-8"))
            data["chunks"] = [
                chunk
                for chunk in data["chunks"]
                if heading_id not in (chunk.get("block_ids") or [])
            ]
            data["chunk_count"] = len(data["chunks"])
            chunks_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            manifest_path = paths[2] / "manifest.json"
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            manifest[0]["chunks"] = len(data["chunks"])
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            report = audit_corpus(*paths, skip_database=True)

        codes = {issue["code"] for issue in report["issues"]}
        self.assertEqual(report["status"], "error")
        self.assertIn("chunk_blocks_missing", codes)

    def test_custom_chunk_size_is_replayed_from_manifest(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            paths = self._build_corpus(Path(temporary_directory), max_chars=24)
            report = audit_corpus(*paths, skip_database=True)

        self.assertEqual(report["status"], "ok", report["issues"])
        self.assertEqual(report["documents"][0]["chunk_max_chars"], 24)
        self.assertEqual(report["documents"][0]["chunk_snapshot_changed"], 0)

    def test_database_chunk_without_embedding_is_reported(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            paths = self._build_corpus(Path(temporary_directory))
            extraction_path = next(
                path for path in paths[1].glob("*.json") if path.name != "manifest.json"
            )
            extraction = json.loads(extraction_path.read_text(encoding="utf-8"))
            chunk_path = next(paths[2].glob("*.chunks.json"))
            chunk = json.loads(chunk_path.read_text(encoding="utf-8"))["chunks"][0]

            document_row = {
                "doc_id": extraction["metadata"]["doc_id"],
                "source_name": extraction["metadata"]["source_name"],
                "index_code": extraction["metadata"].get("index_code"),
                "document_title": extraction["metadata"].get("display_title"),
                "version": extraction["metadata"].get("version"),
                "metadata": extraction["metadata"],
            }
            chunk_row = {
                "chunk_id": chunk["chunk_id"],
                "doc_id": chunk["doc_id"],
                "source_name": chunk["source_name"],
                "index_code": chunk.get("index_code"),
                "document_title": chunk.get("document_title"),
                "version": chunk.get("version"),
                "chunk_type": chunk.get("chunk_type"),
                "citation_label": chunk.get("citation_label"),
                "raw_text": chunk["raw_text"],
                "searchable_text": chunk["searchable_text"],
                "block_ids": chunk.get("block_ids") or [],
                "section_path": chunk.get("section_path") or [],
                "section_labels": chunk.get("section_labels") or [],
                "section_title": chunk.get("section_title"),
                "subsection_title": chunk.get("subsection_title"),
                "item_number": chunk.get("item_number"),
                "heading_number": chunk.get("heading_number"),
                "appendix_number": chunk.get("appendix_number"),
                "appendix_title": chunk.get("appendix_title"),
                "char_count": chunk["char_count"],
                "embedding_model": "bge-m3",
                "metadata": chunk,
                "embedding_missing": True,
                "embedding_dim": None,
            }
            connection = _FakeConnection(document_row, chunk_row)
            issues: list[dict[str, object]] = []
            with (
                patch.object(corpus_audit_module, "database_url", return_value="postgresql://db"),
                patch.object(corpus_audit_module, "connect", return_value=connection),
            ):
                result = corpus_audit_module._audit_database(
                    None,
                    {document_row["doc_id"]: document_row},
                    {chunk_row["chunk_id"]: chunk},
                    issues,
                )

        self.assertTrue(result["checked"])
        self.assertEqual(result["missing_embeddings"], 1)
        self.assertIn("database_embeddings_missing", {issue["code"] for issue in issues})

    def test_database_scalar_retrieval_metadata_mismatch_is_reported(self) -> None:
        expected = {
            "chunk_id": "chunk-1",
            "doc_id": "doc-1",
            "source_name": "document.docx",
            "index_code": "TEST-1.0",
            "document_title": "Политика",
            "version": "1.0",
            "chunk_type": "numbered_item",
            "citation_label": "Политика, пункт 1.1",
            "raw_text": "1.1 Требование",
            "searchable_text": "Пункт: 1.1\n1.1 Требование",
            "block_ids": ["item-1.1"],
            "section_path": ["1"],
            "section_labels": ["1 Общие положения"],
            "section_title": "Общие положения",
            "subsection_title": None,
            "item_number": "1.1",
            "heading_number": None,
            "appendix_number": None,
            "appendix_title": None,
            "char_count": 15,
        }
        document_metadata = {
            "doc_id": "doc-1",
            "source_name": "document.docx",
            "index_code": "TEST-1.0",
            "display_title": "Политика",
            "version": "1.0",
        }
        document_row = {
            "doc_id": "doc-1",
            "source_name": "document.docx",
            "index_code": "TEST-1.0",
            "document_title": "Политика",
            "version": "1.0",
            "metadata": document_metadata,
        }
        chunk_row = {
            **expected,
            "citation_label": "Испорченная цитата",
            "section_labels": ["9 Испорченный раздел"],
            "heading_number": "9",
            "embedding_model": "bge-m3",
            "metadata": expected,
            "embedding_missing": False,
            "embedding_dim": 1024,
        }
        issues: list[dict[str, object]] = []
        with (
            patch.object(corpus_audit_module, "database_url", return_value="postgresql://db"),
            patch.object(
                corpus_audit_module,
                "connect",
                return_value=_FakeConnection(document_row, chunk_row),
            ),
        ):
            corpus_audit_module._audit_database(
                None,
                {"doc-1": {"source_name": "document.docx", "metadata": document_metadata}},
                {"chunk-1": expected},
                issues,
            )

        mismatch = next(
            issue for issue in issues if issue["code"] == "database_chunk_content_mismatch"
        )
        self.assertIn("citation_label", mismatch["details"]["chunks"]["chunk-1"])
        self.assertIn("section_labels", mismatch["details"]["chunks"]["chunk-1"])
        self.assertIn("heading_number", mismatch["details"]["chunks"]["chunk-1"])

    def _build_corpus(
        self,
        root: Path,
        *,
        max_chars: int = 1800,
    ) -> tuple[Path, Path, Path]:
        docs_dir = root / "docs"
        extracted_dir = root / "artifacts" / "extracted"
        chunks_dir = root / "artifacts" / "chunks"
        docs_dir.mkdir(parents=True)
        extracted_dir.mkdir(parents=True)
        chunks_dir.mkdir(parents=True)

        source_path = docs_dir / "Политика.docx"
        document = Document()
        document.add_paragraph("ИНДЕКС НД: TEST-1.0")
        document.add_paragraph("ПОЛИТИКА")
        document.add_paragraph("1 Общие положения")
        document.add_paragraph("1.1 Назначение")
        document.add_paragraph("1.1.1 Первый обязательный пункт документа.")
        document.add_paragraph("1.1.2 Второй обязательный пункт документа.")
        document.save(source_path)

        extracted = extract_docx(source_path)
        extracted_path = extracted_dir / f"{extracted.metadata.doc_id}.json"
        write_extraction(extracted, extracted_path)
        extraction_manifest = [
            {
                "source_name": extracted.metadata.source_name,
                "source_path": extracted.metadata.source_path,
                "doc_id": extracted.metadata.doc_id,
                "output_path": str(extracted_path),
            }
        ]
        (extracted_dir / "manifest.json").write_text(
            json.dumps(extraction_manifest, ensure_ascii=False),
            encoding="utf-8",
        )

        chunked = chunk_document(extracted.to_dict(), max_chars=max_chars)
        chunks_path = chunks_dir / f"{extracted.metadata.doc_id}.chunks.json"
        chunks_path.write_text(json.dumps(chunked, ensure_ascii=False), encoding="utf-8")
        chunk_manifest = [
            {
                "doc_id": extracted.metadata.doc_id,
                "source_name": extracted.metadata.source_name,
                "extracted_path": str(extracted_path),
                "chunked_path": str(chunks_path),
                "chunks": chunked["chunk_count"],
                "max_chars": max_chars,
            }
        ]
        (chunks_dir / "manifest.json").write_text(
            json.dumps(chunk_manifest, ensure_ascii=False),
            encoding="utf-8",
        )
        return docs_dir, extracted_dir, chunks_dir


def _inject_footnote(path: Path, text: str, *, add_reference: bool = True) -> None:
    with ZipFile(path) as source:
        files = {item.filename: source.read(item.filename) for item in source.infolist()}

    content_types = files["[Content_Types].xml"].decode("utf-8")
    content_types = content_types.replace(
        "</Types>",
        (
            '<Override PartName="/word/footnotes.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
            "</Types>"
        ),
    )
    files["[Content_Types].xml"] = content_types.encode("utf-8")

    relationships_path = "word/_rels/document.xml.rels"
    relationships = files[relationships_path].decode("utf-8")
    relationships = relationships.replace(
        "</Relationships>",
        (
            '<Relationship Id="rIdAuditFootnotes" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" '
            'Target="footnotes.xml"/>'
            "</Relationships>"
        ),
    )
    files[relationships_path] = relationships.encode("utf-8")

    if add_reference:
        document_xml = files["word/document.xml"].decode("utf-8")
        reference = '<w:r><w:footnoteReference w:id="1"/></w:r>'
        document_xml = document_xml.replace("</w:p>", f"{reference}</w:p>", 1)
        files["word/document.xml"] = document_xml.encode("utf-8")

    escaped = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )
    files["word/footnotes.xml"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:footnote w:id="1"><w:p><w:r><w:t>'
        f"{escaped}"
        "</w:t></w:r></w:p></w:footnote>"
        "</w:footnotes>"
    ).encode("utf-8")

    temporary_path = path.with_suffix(".rewrite.docx")
    with ZipFile(temporary_path, "w", ZIP_DEFLATED) as destination:
        for name, payload in files.items():
            destination.writestr(name, payload)
    temporary_path.replace(path)


class _FakeRows:
    def __init__(self, rows: list[dict[str, object]]) -> None:
        self.rows = rows

    def fetchall(self) -> list[dict[str, object]]:
        return self.rows


class _FakeConnection:
    def __init__(self, document: dict[str, object], chunk: dict[str, object]) -> None:
        self.document = document
        self.chunk = chunk

    def __enter__(self) -> _FakeConnection:
        return self

    def __exit__(self, *args: object) -> None:
        return None

    def execute(self, query: str) -> _FakeRows:
        normalized = " ".join(query.split())
        if "FROM doc_documents" in normalized:
            return _FakeRows([self.document])
        if "FROM doc_chunks" in normalized:
            return _FakeRows([self.chunk])
        if "FROM pg_indexes" in normalized:
            return _FakeRows([{"indexname": "doc_chunks_embedding_hnsw"}])
        raise AssertionError(f"Unexpected query: {normalized}")


if __name__ == "__main__":
    unittest.main()
